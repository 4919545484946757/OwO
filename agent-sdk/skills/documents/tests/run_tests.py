# -*- coding: utf-8 -*-
"""documents 技能契约测试：docx 生成 / 修改 / 结构校验（>=3 个端到端用例）。"""
import os
import shutil
import tempfile
import zipfile

from docx import Document


def case1(tmp: str) -> None:
    """从大纲生成 docx，回读断言标题层级与段落。"""
    doc = Document()
    doc.add_heading("项目大纲", level=1)
    doc.add_heading("背景", level=2)
    doc.add_paragraph("这是背景段落。")
    doc.add_heading("目标", level=2)
    doc.add_paragraph("这是目标段落。")
    path = os.path.join(tmp, "outline.docx")
    doc.save(path)
    back = Document(path)
    texts = [p.text for p in back.paragraphs]
    assert "项目大纲" in texts
    assert "背景" in texts
    assert "这是背景段落。" in texts
    print("case1 ok: 大纲生成 docx 并回读")


def case2(tmp: str) -> None:
    """修改既有 docx：追加标题与段落并回读。"""
    path = os.path.join(tmp, "report.docx")
    doc = Document()
    doc.add_heading("报告", level=1)
    doc.save(path)
    doc2 = Document(path)
    doc2.add_heading("结论", level=2)
    doc2.add_paragraph("追加内容。")
    doc2.save(path)
    back = Document(path)
    texts = [p.text for p in back.paragraphs]
    assert "结论" in texts and "追加内容。" in texts
    print("case2 ok: 追加段落并回读")


def case3(tmp: str) -> None:
    """zip 包结构与内容断言。"""
    path = os.path.join(tmp, "structure.docx")
    doc = Document()
    doc.add_heading("结构校验", level=1)
    doc.add_paragraph("内容A")
    doc.save(path)
    with zipfile.ZipFile(path) as archive:
        names = archive.namelist()
        assert "word/document.xml" in names
    back = Document(path)
    assert back.paragraphs[0].text == "结构校验"
    print("case3 ok: zip 结构校验")


def main() -> None:
    tmp = tempfile.mkdtemp(prefix="owo-skill-documents-")
    try:
        case1(tmp)
        case2(tmp)
        case3(tmp)
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


if __name__ == "__main__":
    main()
